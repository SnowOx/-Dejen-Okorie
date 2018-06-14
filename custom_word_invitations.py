#! python3
# custom_word_invitations.py

# Remember that you must define a custom style in a document before py docx can use it

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

guests = ['Prof Plum', 'Miss Scarlet', 'Col Mustard', 'Al Sweigart', 'Robocop',
          ' Abbot', 'Abbott', 'Abby', 'Abdel', 'Abdiel', 'Abe', 'Abednego', 'Abejundio',
          'Abel', 'Abelard', 'Abenster', 'Abenzio', 'Abercrombie', 'Aberdeen', 'Abhay',
          'Abhi', 'Abhijit', 'Abhinav', 'Abhishek', 'Abie', 'Abijah', 'Abner', 'Abrasha',
          'Absolut', 'Acacio', 'Acapulco', 'Accord', 'Ace', 'Acer', 'Achar', 'Achilles',
          'Achimus', 'Achitu', 'Achyuta', 'Ackbar', 'Ackley', 'Acrobat', 'Action', 'Ad lib',
          'Adalardo', 'Adalgiso', 'Adamore', 'Addai', 'Addison', 'Addo', 'Adelbert', 'Adelfried',
          'Adelino', 'Adelmo', 'Adelphos', 'Ademaro', 'Aditya', 'Adlai', 'Adler', 'Adley', 'Admes',
          'Admira', 'Admiral', 'Admiral Kirk', 'Admon', 'Adolph', 'Adon', 'Adonis', 'Adrastos', 'Adrian',
          'Aeneas', 'Aero', 'Aesop']

def write_invite_text():
    for guest in guests:
        document.add_paragraph()
        para_1 = document.add_paragraph()
        para_1.add_run('It would be a pleasure to have the company of')
        para_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_1.style = styles['Normal']
        document.add_paragraph()
        para_2 = document.add_paragraph(guest)
        para_2.style = styles['TD Script']
        para_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()
        para_3 = document.add_paragraph()
        para_3.style = styles['Normal']
        para_3.add_run('at 1101 Memory Lane on the Evening of')
        para_3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()
        para_4 = document.add_paragraph('April 1st')
        para_4.style = styles['Normal']
        para_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()
        para_5 = document.add_paragraph('at 7 o\'clock')
        para_5.style = styles['Normal']
        para_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

document = Document('tddemo.docx')
styles = document.styles

write_invite_text()
document.save('Completed_Invitations.docx')
